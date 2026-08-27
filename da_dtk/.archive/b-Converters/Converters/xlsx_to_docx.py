# -*- coding: utf-8 -*-
"""
This script reads a structured XLSX (Excel) file and converts it into a 
DOCX (Word) file containing a properly formatted table.

This EFFICIENT version pre-builds the table structure to handle very large
Excel files much faster, avoiding performance bottlenecks.

Required libraries: pip install openpyxl python-docx
"""

import sys

# --- Library Installation Checks ---
try:
    import openpyxl
except ImportError:
    print("Error: The 'openpyxl' library is required.")
    print("Please install it by running: pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Error: The 'python-docx' library is required.")
    print("Please install it by running: pip install python-docx")
    sys.exit(1)

def convert_xlsx_to_docx(xlsx_file_path, docx_file_path):
    """
    Reads an XLSX file and efficiently converts the first sheet into a table 
    in a DOCX document, optimized for large files.
    """
    try:
        print("Loading the Excel workbook into memory...")
        workbook = openpyxl.load_workbook(xlsx_file_path)
        sheet = workbook.active
    except FileNotFoundError:
        print(f"Error: The file '{xlsx_file_path}' was not found.")
        return
    except Exception as e:
        print(f"An error occurred while opening the Excel file: {e}")
        return

    # --- Create a new Word Document ---
    document = Document()
    document.add_heading(f'Data from {xlsx_file_path}', level=1)
    document.add_paragraph('This table contains the data converted from the Excel spreadsheet.')

    # --- EFFICIENT TABLE CREATION ---
    # 1. Get the total number of rows and columns from the Excel sheet.
    total_rows = sheet.max_row
    total_cols = sheet.max_column
    
    print(f"Found {total_rows} rows and {total_cols} columns.")
    print("Pre-building the table structure in the Word document...")

    # 2. Create the table with the exact dimensions needed. This is much faster.
    table = document.add_table(rows=total_rows, cols=total_cols)
    table.style = 'Table Grid' # Apply a visible grid style

    print("Populating the table with data...")

    # 3. Iterate through the sheet and populate the pre-made cells.
    for i, row in enumerate(sheet.iter_rows()):
        # Get the corresponding row in the Word table
        table_row_cells = table.rows[i].cells
        for j, cell in enumerate(row):
            # Get the corresponding cell and fill it
            # Ensure None values are written as empty strings
            table_row_cells[j].text = str(cell.value if cell.value is not None else "")
            
            # Make the header row (i==0) bold
            if i == 0:
                table_row_cells[j].paragraphs[0].runs[0].bold = True

    # --- Save the Document ---
    try:
        print("Saving the final Word document...")
        document.save(docx_file_path)
        print(f"✅ Successfully converted '{xlsx_file_path}' to '{docx_file_path}'")
    except Exception as e:
        print(f"An error occurred while saving the Word document: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python xlsx_to_docx_converter_efficient.py <input_xlsx_file> <output_docx_file>")
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    convert_xlsx_to_docx(input_file, output_file)